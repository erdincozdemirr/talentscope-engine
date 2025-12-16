# -*- coding: utf-8 -*-
from pathlib import Path
from typing import List


def extract_text_from_pdf(path: Path) -> str:
    from pdfminer.high_level import extract_text
    try:
        return extract_text(str(path)) or ""
    except Exception:
        return ""


def extract_text_from_docx(path: Path) -> str:
    from docx import Document
    try:
        doc = Document(str(path))
        parts: List[str] = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    txt = (cell.text or "").strip()
                    if txt:
                        parts.append(txt)
        return "\n".join(parts)
    except Exception:
        return ""



def extract_text_from_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def extract_file_content(file_path: str) -> str:
    p = Path(file_path)
    suf = p.suffix.lower()
    if suf == ".pdf":
        return extract_text_from_pdf(p)
    if suf == ".docx":
        return extract_text_from_docx(p)
    if suf == ".txt":
        return extract_text_from_txt(p)
    raise ValueError(f"Unsupported file type: {suf}. Only .pdf, .docx, and .txt are supported.")


def extract_resume_text(file_path: str) -> str:
    """Wrapper for backward compatibility, though extract_file_content covers all."""
    return extract_file_content(file_path)

